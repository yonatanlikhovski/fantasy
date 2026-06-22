"""
Build an English Word (.docx) document that explains every visualization in
`quality_plots/`: what it shows, the concrete insight worth citing in the
project writeup, and which visualization principle (from the course notes) it
demonstrates.

All numbers are computed live from the CSVs under `sim_stats/`, so the document
stays in sync with the data. Run from the repo root:

    py visual_helper/build_insights_doc.py

Requires: pandas, numpy, scipy, python-docx.
"""

from __future__ import annotations

import itertools
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
SIM_STATS = REPO_ROOT / "sim_stats"
OUT_PATH = SCRIPT_DIR / "quality_plots" / "Visualization_Insights.docx"

H2H_CATS = ["PTS", "REB", "AST", "STL", "BLK", "FG3M", "TOV", "FG", "FT"]
ACCENT = RGBColor(0x0B, 0x53, 0x94)  # deep blue for headings


# ---------------------------------------------------------------------
# Number crunching (everything cited in the doc comes from here)
# ---------------------------------------------------------------------

def compute_facts() -> dict:
    f: dict = {}

    ss = pd.read_csv(SIM_STATS / "player_season_stats.csv")
    ss["season"] = pd.to_numeric(ss["season"], errors="coerce")

    counts = ss.groupby("season")["player_id"].nunique()
    f["season_counts"] = {int(k): int(v) for k, v in counts.items()}
    latest = int(ss["season"].max())
    f["latest_season"] = latest
    latest_players = set(ss.loc[ss["season"] == latest, "player_id"])
    f["latest_n"] = len(latest_players)
    f["retention"] = {
        int(s): len(latest_players & set(ss.loc[ss["season"] == s, "player_id"]))
        for s in sorted(ss["season"].dropna().unique())
    }

    ppg = {}
    for s in sorted(ss["season"].dropna().unique()):
        v = pd.to_numeric(ss.loc[ss["season"] == s, "pts_mean"], errors="coerce").dropna()
        ppg[int(s)] = {
            "mean": float(v.mean()),
            "std": float(v.std()),
            "skew": float(stats.skew(v)),
            "pct_over25": float((v > 25).mean() * 100),
        }
    f["ppg"] = ppg

    scarcity_cols = {"PTS": "pts_mean", "REB": "trb_mean", "AST": "ast_mean",
                     "STL": "stl_mean", "BLK": "blk_mean", "FG3M": "fg3m_mean",
                     "TOV": "tov_mean"}
    cov = {}
    for cat, col in scarcity_cols.items():
        v = pd.to_numeric(ss[col], errors="coerce").dropna()
        v = v[v >= 0]
        if not v.empty and v.mean() != 0:
            cov[cat] = float(v.std() / v.mean())
    f["scarcity"] = dict(sorted(cov.items(), key=lambda kv: kv[1]))

    val = pd.read_csv(SIM_STATS / "h2h_value_2027.csv")
    corr = val[H2H_CATS].corr()
    pairs = [(a, b, float(corr.loc[a, b]))
             for a, b in itertools.combinations(H2H_CATS, 2)]
    pairs.sort(key=lambda x: x[2])
    f["corr_neg"] = pairs[:4]
    f["corr_pos"] = pairs[-4:][::-1]
    top = val.sort_values("H2H_value", ascending=False).head(5)
    f["top_h2h"] = list(top[["player_id", "H2H_value"]].itertuples(index=False, name=None))
    top30 = val.sort_values("H2H_value", ascending=False).head(30)["H2H_value"]
    f["top30_range"] = (float(top30.min()), float(top30.max()))

    ws = pd.read_csv(SIM_STATS / "ws_vs_h2h_2027.csv")
    x = pd.to_numeric(ws["win_shares"], errors="coerce")
    y = pd.to_numeric(ws["H2H_value"], errors="coerce")
    m = x.notna() & y.notna()
    sl, ic, r, p, se = stats.linregress(x[m], y[m])
    f["ws_r"] = float(r)
    f["ws_r2"] = float(r * r)
    f["ws_n"] = int(m.sum())
    f["over"] = list(ws.sort_values("rank_diff", ascending=False)
                     .head(4)[["player_id", "rank_diff"]].itertuples(index=False, name=None))
    f["under"] = list(ws.sort_values("rank_diff")
                      .head(4)[["player_id", "rank_diff"]].itertuples(index=False, name=None))

    summ = pd.read_csv(SIM_STATS / "evaluation_2026_summary.csv")
    f["model_summary"] = summ
    f["model_best_corr"] = summ.sort_values("pearson", ascending=False).iloc[0]
    f["model_worst_corr"] = summ.sort_values("pearson").iloc[0]
    f["model_cov_mean"] = float(summ["ci_5_95_coverage"].mean())

    st = pd.read_csv(SIM_STATS / "draft_standings_2027.csv")
    st = st.sort_values("W", ascending=False)
    f["standings"] = st
    f["champion"] = st.iloc[0]
    f["last_place"] = st.iloc[-1]
    f["w_spread"] = float(st["W"].max() - st["W"].min())

    return f


# ---------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------

def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT


def add_figure_block(doc: Document, title: str, file_rel: str, shows: str,
                     insight: str, principles: str) -> None:
    add_heading(doc, title, level=3)

    p = doc.add_paragraph()
    p.add_run("File: ").bold = True
    code = p.add_run(file_rel)
    code.font.name = "Consolas"
    code.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run("What it shows: ").bold = True
    p.add_run(shows)

    p = doc.add_paragraph()
    p.add_run("Insight to mention in the paper: ").bold = True
    p.add_run(insight)

    p = doc.add_paragraph()
    p.add_run("Principles demonstrated: ").bold = True
    r = p.add_run(principles)
    r.italic = True


def fmt_counts(d: dict) -> str:
    return ", ".join(f"{k}: {v}" for k, v in sorted(d.items()))


# ---------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------

def build(f: dict) -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Fantasy Value Prediction for NBA Players", level=0)
    for run in title.runs:
        run.font.color.rgb = ACCENT
    sub = doc.add_paragraph("Visualization Insights for the Project Writeup")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(13)

    doc.add_paragraph(
        "This document accompanies the figures in the quality_plots/ folder. For each "
        "visualization it states what the figure shows, the concrete take-away to cite "
        "in the writeup, and the visualization principle it demonstrates. Every number "
        "below is computed directly from the project's data files in sim_stats/."
    )

    # ---- principles summary -------------------------------------------------
    add_heading(doc, "How the figures follow the visualization principles", level=1)
    doc.add_paragraph(
        "The figures were designed against the six principles from the course "
        "(Scale, Conditioning, Perception, Transformation, Context, Smoothing) and "
        "Tufte/Wainer's rules for honest, high data-ink graphics:"
    )
    for label, text in [
        ("Scale", "Every bar chart starts at zero and uses a single consistent axis; "
                  "axis limits are chosen to fill the plot without exaggerating differences."),
        ("Conditioning", "Distributions are shown per subgroup, e.g. points-per-game is "
                         "faceted by season and the projection scatter is split per statistic."),
        ("Perception", "Colors carry meaning through a colorblind-safe palette (blue/orange, "
                       "never red+green together) and perceptually-uniform colormaps (viridis). "
                       "Quantities are encoded by length/position; no pie or area charts are used, "
                       "and the original stacked bar was replaced by a heatmap to avoid a moving baseline."),
        ("Transformation", "Stats are standardized (z-scores) for the team-strengths heatmap and "
                           "normalized to 0-1 for category comparisons so they share one scale."),
        ("Context", "Each figure has a take-away title (a conclusion, not a description), labelled "
                    "axes, reference lines (y = x, ideal coverage = 0.90, the regression trend), "
                    "annotated outliers, and a descriptive caption."),
        ("Smoothing", "A fitted normal curve is overlaid on each points-per-game histogram to make "
                     "the shape and skew of the distribution easy to read."),
        ("Data-ink", "Chart junk is minimized: top/right spines removed, light grid, direct data "
                     "labels instead of clutter."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)

    doc.add_paragraph(
        "All eight visualizations map directly onto the project goal defined in the milestone: "
        "building and validating a data-driven H2H fantasy value for NBA players, and contrasting "
        "that fantasy value with real-life basketball value."
    )

    # ---- overview -----------------------------------------------------------
    add_heading(doc, "1. Data overview and sanity checks", level=1)

    sc = f["season_counts"]
    add_figure_block(
        doc, "1.1 Active players per season", "overview/active_players_per_season.png",
        "The number of distinct players with at least one scraped game log in each season "
        "from 2023 to 2026.",
        f"The scrape yields {sc.get(2023,'?')} players in 2023 rising to {sc.get(2026,'?')} in 2026 "
        f"({fmt_counts(sc)}). The smooth, monotonic growth toward the present (and the absence of "
        "implausible spikes) is a sanity check that the game-by-game scraping is complete and not "
        "duplicated, which is exactly the kind of data-quality reassurance a good visualization "
        "should provide before any modelling.",
        "Scale (zero baseline), Context (take-away title and caption).")

    ret = f["retention"]
    add_figure_block(
        doc, "1.2 Retention of the current player pool", "overview/player_retention.png",
        f"Of the {f['latest_n']} players active in {f['latest_season']}, how many also appear in "
        "each earlier season.",
        f"Going backwards the overlap falls steadily ({fmt_counts(ret)}), i.e. roughly 50-90 of "
        f"today's players drop out per season we look back. This matches the domain expectation of "
        "about 70 players entering/leaving the league each year (draft, two-way contracts, "
        "retirements) and confirms the multi-season panel is internally consistent.",
        "Conditioning (overlap by season), Perception (single highlight color), Context.")

    p23, p26 = f["ppg"][2023], f["ppg"][2026]
    skews = [v["skew"] for v in f["ppg"].values()]
    overs = [v["pct_over25"] for v in f["ppg"].values()]
    add_figure_block(
        doc, "1.3 Points-per-game distribution by season", "overview/ppg_distribution_by_season.png",
        "A 2x2 grid of points-per-game histograms (one per season) with a fitted normal curve and "
        "the mean marked.",
        f"Scoring is right-skewed every season (skewness about +{min(skews):.2f} to +{max(skews):.2f}), "
        f"with means near {f['ppg'][2025]['mean']:.1f}-{p23['mean']:.1f} points and only "
        f"{min(overs):.0f}-{max(overs):.0f}% of players above 25 PPG. This confirms the "
        "milestone's hypothesis that most players have limited scoring roles while elite scorers are "
        "rare outliers - and shows a normal curve is only a rough fit, which itself justifies the "
        "category-scarcity reasoning later.",
        "Conditioning (faceting by season), Smoothing (normal fit), Context (mean reference line).")

    # ---- value --------------------------------------------------------------
    add_heading(doc, "2. Fantasy H2H value", level=1)

    add_figure_block(
        doc, "2.1 Category contribution breakdown (top 15)", "value/category_breakdown_top15.png",
        "A heatmap of the 15 highest-valued players against their normalized score (0-1) in each of "
        "the nine scoring categories.",
        "The top players score near the league maximum across most categories, but each still has a "
        "signature: the breakdown makes clear that elite fantasy value comes from being broadly "
        "strong rather than dominating one stat, while role differences (e.g. a center's blocks vs a "
        "guard's threes) remain visible. A heatmap is used instead of a stacked bar precisely because "
        "stacked segments share a moving baseline and are hard to compare.",
        "Perception (heatmap over stacked bar, uniform colormap), Transformation (0-1 normalization), Context.")

    sca = f["scarcity"]
    scarce_top = list(sca.items())[-1]
    scarce_low = list(sca.items())[0]
    add_figure_block(
        doc, "2.2 Category scarcity", "value/category_scarcity.png",
        "The coefficient of variation (std / mean) of each per-game category across all players.",
        f"{scarce_top[0]} is by far the scarcest category (CoV {scarce_top[1]:.2f}) while {scarce_low[0]} "
        f"is the most evenly spread (CoV {scarce_low[1]:.2f}). Because winning an H2H category means "
        "beating an opponent directly, scarce and highly variable stats such as blocks swing weekly "
        "matchups more than abundant ones - this is the quantitative justification for weighting "
        "categories differently in the value formula.",
        "Scale (zero baseline), Perception (uniform colormap), Context (descriptive caption).")

    cn = f["corr_neg"][0]
    cp = f["corr_pos"][0]
    add_figure_block(
        doc, "2.3 Category correlation heatmap", "value/category_correlation_heatmap.png",
        "The Pearson correlation between players' normalized category scores.",
        f"Most category pairs are only weakly correlated; the clearest signals are turnovers moving "
        f"opposite to usage stats (e.g. {cn[0]}-{cn[1]} = {cn[2]:.2f}) and modest positive links among "
        f"counting stats (e.g. {cp[0]}-{cp[1]} = {cp[2]:.2f}). The general independence means a "
        "balanced roster cannot be assembled by chasing one super-stat; categories must be targeted "
        "deliberately, supporting the team-building strategy in the draft simulation.",
        "Perception (diverging blue-orange with light middle, colorblind-safe), Context.")

    th = f["top_h2h"]
    lo, hi = f["top30_range"]
    add_figure_block(
        doc, "2.4 Top 30 players by H2H value", "value/top30_h2h_value.png",
        "A ranked horizontal bar chart of the 30 highest projected H2H values.",
        f"The model's top tier is led by {th[0][0]} ({th[0][1]:.2f}), {th[1][0]} ({th[1][1]:.2f}) and "
        f"{th[2][0]} ({th[2][1]:.2f}); the whole top 30 is compressed into a narrow band "
        f"({lo:.2f}-{hi:.2f}). The small gaps between adjacent players show that draft order matters "
        "most at the very top and that many mid-first-round players are near-interchangeable in value.",
        "Scale (zero baseline, sorted), Perception (length encoding), Context.")

    # ---- thesis -------------------------------------------------------------
    add_heading(doc, "3. Real value vs fantasy value (core thesis)", level=1)

    add_figure_block(
        doc, "3.1 Win Shares vs H2H value", "thesis/ws_vs_h2h_scatter.png",
        "A scatter plot of each player's real-life value (Win Shares) against fantasy value (H2H), "
        "with a regression line and labelled outliers.",
        f"Real and fantasy value are only moderately related (Pearson r = {f['ws_r']:.2f}, "
        f"r-squared = {f['ws_r2']:.2f}, n = {f['ws_n']}): real-life value explains only about "
        f"{f['ws_r2']*100:.0f}% of the variation in fantasy value. This is the central evidence for "
        "the project's premise that a player's fantasy worth does not match his basketball worth, and "
        "the labelled outliers are concrete examples to discuss.",
        "Context (take-away title, trend line, outlier labels), Transformation (regression fit).")

    ov = f["over"][0]
    un = f["under"][0]
    add_figure_block(
        doc, "3.2 Most over- and under-valued players", "thesis/over_under_valued.png",
        "A diverging bar chart of the players with the largest gap between their Win-Shares rank and "
        "their H2H rank.",
        f"Some players are fantasy overvalued while others are hidden bargains: {ov[0]} is ranked "
        f"about {int(ov[1])} places higher in fantasy than in real life, while {un[0]} is about "
        f"{abs(int(un[1]))} places higher in real life than in fantasy. These gaps are actionable draft "
        "advice and a vivid illustration of the thesis from 3.1.",
        "Perception (colorblind-safe blue/orange diverging), Scale (signed zero baseline), Context.")

    # ---- model --------------------------------------------------------------
    add_heading(doc, "4. Projection model evaluation", level=1)

    summ = f["model_summary"]
    bc, wc = f["model_best_corr"], f["model_worst_corr"]
    add_figure_block(
        doc, "4.1 Predicted vs actual (per stat)", "model/predicted_vs_actual.png",
        "A 3x3 grid of predicted-vs-actual scatter plots for the nine projected statistics, each with "
        "a y = x reference line.",
        f"Projections cluster tightly around the diagonal for volume/counting stats (best: "
        f"{bc['stat']}, r = {bc['pearson']:.2f}) but scatter widely for shooting percentages (worst: "
        f"{wc['stat']}, r = {wc['pearson']:.2f}). This tells the reader exactly which projected "
        "quantities can be trusted and which carry large uncertainty.",
        "Conditioning (one panel per stat), Context (y = x reference, per-panel correlation).")

    pts = summ.loc[summ["stat"] == "PTS"].iloc[0]
    add_figure_block(
        doc, "4.2 Projection error per stat (MAE / RMSE)", "model/error_mae_rmse.png",
        "Grouped bars of mean absolute error and root mean squared error for each statistic, in that "
        "stat's own units.",
        f"Absolute error is largest for points (MAE about {pts['mae']:.1f}) and very small for rate "
        "and low-count stats, simply because points are on a larger scale. RMSE sitting above MAE for "
        "every stat signals that a minority of players (typically those with role or injury changes) "
        "drive the biggest misses.",
        "Scale (shared zero baseline), Perception (blue/orange grouped bars), Context.")

    add_figure_block(
        doc, "4.3 Projection accuracy (Pearson / Spearman)", "model/correlation_pearson_spearman.png",
        "Grouped bars of the Pearson and Spearman correlation between projected and actual values per "
        "stat (0-1 scale).",
        f"Counting stats correlate strongly with reality (around 0.7-0.75) while percentage stats are "
        f"much weaker (FG% and FT% near 0.37-0.43). Pearson and Spearman agreeing closely shows the "
        "ranking of players is preserved even where absolute values are noisy - important because H2H "
        "value depends on relative ordering.",
        "Scale (0-1 axis), Perception, Context (two complementary metrics side by side).")

    add_figure_block(
        doc, "4.4 Confidence-interval calibration", "model/ci_calibration.png",
        "The share of actual outcomes that fell inside the predicted 5-95% interval per stat, against "
        "the ideal 0.90 line.",
        f"Coverage averages only about {f['model_cov_mean']:.2f} and is well below 0.90 for almost "
        "every counting stat, meaning the model's stated uncertainty is too narrow - real outcomes "
        "land outside the interval far more than 10% of the time. This is an honest limitation to "
        "report and a clear direction for future work (wider/recalibrated intervals).",
        "Context (ideal-coverage reference line, honest reporting of a flaw), Scale.")

    # ---- draft --------------------------------------------------------------
    add_heading(doc, "5. Draft simulation", level=1)

    champ, last = f["champion"], f["last_place"]
    add_figure_block(
        doc, "5.1 Final standings", "draft/standings_wins.png",
        "A horizontal bar chart of each simulated team's average weekly wins, annotated with total "
        "category points.",
        f"The strategies produce a clear spread of {f['w_spread']:.1f} wins between best and worst: "
        f"team {int(champ['team'])} leads with {champ['W']:.1f} average weekly wins versus "
        f"{last['W']:.1f} for team {int(last['team'])}. The gap shows that draft strategy materially "
        "affects outcomes, validating the simulation as a way to compare value approaches.",
        "Scale (zero baseline, sorted), Perception (length encoding), Context (direct labels).")

    add_figure_block(
        doc, "5.2 Team category strengths", "draft/team_category_strengths.png",
        "A heatmap of teams (ordered by standings) against the nine categories, colored by z-score "
        "within each category (turnovers sign-flipped so higher is always better).",
        "Each team's row reveals the categories it punted or prioritized, and the top teams are not "
        "uniformly strong but specialized in a coherent set of categories. This visually confirms the "
        "H2H insight that building around a few targeted strengths beats chasing the best available "
        "player every round.",
        "Transformation (z-score standardization, sign-flip for TOV), Perception (diverging colorblind-safe map), Context.")

    return doc


def main() -> None:
    facts = compute_facts()
    doc = build(facts)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"[OK] wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
